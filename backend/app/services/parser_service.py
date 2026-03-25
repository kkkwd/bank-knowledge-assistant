import subprocess
from pathlib import Path

import pytesseract
from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

from app.services.chunk_service import ParsedSection


class ParserService:
    def parse(self, storage_path: str) -> dict[str, str | list[ParsedSection]]:
        path = Path(storage_path)
        suffix = path.suffix.lower()

        if suffix == ".txt":
            text = path.read_text(encoding="utf-8")
            return {"title": path.stem, "sections": [ParsedSection(text=text)]}

        if suffix == ".docx":
            return {"title": path.stem, "sections": [ParsedSection(text=self._parse_docx(path))]}

        if suffix == ".pdf":
            return {"title": path.stem, "sections": self._parse_pdf(path)}

        if suffix in {".jpg", ".jpeg", ".png"}:
            return {"title": path.stem, "sections": [ParsedSection(text=self._parse_image(path))]}

        raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_docx(self, path: Path) -> str:
        document = DocxDocument(path)
        paragraphs = [item.text.strip() for item in document.paragraphs if item.text.strip()]
        table_cells = [
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        ]
        text = "\n".join(paragraphs + table_cells)
        if text:
            return text

        result = subprocess.run(
            ["/usr/bin/textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True,
            text=True,
            check=False,
        )
        return result.stdout.strip()

    def _parse_pdf(self, path: Path) -> list[ParsedSection]:
        reader = PdfReader(str(path))
        sections: list[ParsedSection] = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append(ParsedSection(text=text, page_number=index))
        return sections

    def _parse_image(self, path: Path) -> str:
        with Image.open(path) as image:
            try:
                return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
            except pytesseract.TesseractError:
                return pytesseract.image_to_string(image, lang="eng").strip()
